"""DOCX 结构化切块：移植 RAGFlow Manual 算法（标题层级栈 + 表格原子化）。

与 ingestion.parse_files 的默认 SentenceSplitter 策略并存，由 settings.chunking.strategy
= "ragflow_manual_docx" 开启。零新依赖（仅 python-docx + SentenceSplitter 兜底）。

设计依据：
  - MultiDocFusion (arXiv:2604.12352) 证明「层级感知切块」比纯文本切块 retrieval
    precision +8–15%。其核心 = DSHP-LLM 重建章节父子树 → DFS 累积祖先标题路径切块。
  - RAGFlow Manual (rag/app/manual.py) 已把这套思想工程化：question_stack 维护父子
    标题栈 = 论文的 Header Tree；表格 tbls 独立处理 = 论文的表格节点。
  - 本模块移植 RAGFlow Manual 的工程实现。DOCX 自带 Heading N 样式，层级免费，
    无需论文里给扫描件用的 DSHP-LLM 视觉推断，也不加载任何模型。
"""
from __future__ import annotations

import hashlib
import logging
from pathlib import Path

from core.rag.llamaindex.file_routing import serialize_table

logger = logging.getLogger(__name__)

# OOXML 命名空间
_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_DML = "http://schemas.openxmlformats.org/drawingml/2006/main"
_PIC = "http://schemas.openxmlformats.org/drawingml/2006/picture"
_RELS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

# pStyle val → Heading 层级。覆盖英文 style_id（Heading1..6）与中文 Word 常见的
# OOXML 内置数字 id（"1".."6"）。与 file_routing._H1_SVALS 同源，扩展到 H1-H6。
_HEADING_LEVELS: dict[str, int] = {f"Heading{i}": i for i in range(1, 7)}
for _i in range(1, 7):
    _HEADING_LEVELS[str(_i)] = _i


def _heading_level(p_element) -> int:
    """从 <w:p> 的 pStyle 判定 Heading 层级；非标题返回 0。"""
    pStyle = p_element.find(f".//{{{_NS}}}pStyle")
    if pStyle is None:
        return 0
    sval = pStyle.get(f"{{{_NS}}}val") or ""
    return _HEADING_LEVELS.get(sval, 0)


def _paragraph_text(p_element) -> str:
    """提取 <w:p> 的纯文本。"""
    return "".join(t.text or "" for t in p_element.iter(f"{{{_NS}}}t")).strip()


def _has_image(p_element) -> bool:
    """检测 <w:p> 是否含图片（blip / pic）。"""
    return (
        p_element.find(f".//{{{_DML}}}blip") is not None
        or p_element.find(f".//{{{_PIC}}}pic") is not None
    )


def _image_blob_sha(p_element, doc) -> str | None:
    """取段落首张嵌入图的「原图 blob sha256」。

    经 <a:blip r:embed="rIdN"> → doc.part.rels[rId].target_part.blob → sha256。
    这个 blob 与 image_extractor._extract_docx_images 读的是同一份 target_part.blob，
    故两侧 sha 同口径，DOCX 占位符 [[IMG:sha16]] 能与位置清单 image_desc_by_blob 精确 join。
    无图 / 取不到 rels → None。
    """
    blip = p_element.find(f".//{{{_DML}}}blip")
    if blip is None:
        return None
    rId = blip.get(f"{{{_RELS}}}embed")
    if not rId:
        return None
    try:
        return hashlib.sha256(doc.part.rels[rId].target_part.blob).hexdigest()
    except (KeyError, AttributeError):
        return None


def _split_oversize(text: str, max_chars: int) -> list[str]:
    """超长正文块回退 SentenceSplitter 二次切（表格块不走这里，仍原子）。"""
    try:
        from llama_index.core.node_parser import SentenceSplitter
        from llama_index.core.schema import Document as LlamaDocument
    except ImportError:
        # ingestion 进程必装 LlamaIndex（顶部强制 import）；此分支仅为模块独立可测
        step = max(1, max_chars - max(1, max_chars // 10))
        return [text[i:i + max_chars] for i in range(0, len(text), step)] or [text]

    overlap = max(1, max_chars // 10)
    nodes = SentenceSplitter(
        chunk_size=max_chars, chunk_overlap=overlap, tokenizer=lambda t: t,
    ).get_nodes_from_documents([LlamaDocument(text=text)])
    out = [n.get_content().strip() for n in nodes if n.get_content().strip()]
    return out or [text.strip()]


def chunk_docx_structured(
    file_path: str,
    *,
    max_section_chars: int,
) -> tuple[list[str], list[str]]:
    """移植 RAGFlow Manual 的 DOCX 结构化切块。

    Args:
        file_path: .docx 文件路径。
        max_section_chars: 单个文本块正文的字符上限；超过则对该块内部文本回退
            SentenceSplitter 二次切（表格块不受影响，仍原子）。由调用方传
            LLAMA_INDEX_CHUNK_SIZE，与默认切块器块大小对齐。

    Returns:
        (chunks, sections)，等长配对：
          chunks[i]  = 第 i 块正文（表格块 = 整表序列化文本；文本块 = 祖先标题下
                       累积的正文）。不含【章节:】前缀——前缀由 ingestion 统一注入。
          sections[i]= 第 i 块的祖先标题路径（如 "实验三 > 串联电路"），供注入【章节:】。
        表格块与文本块按文档原始顺序交错。

    两招（移植自 RAGFlow rag/app/manual.py Docx + rag/nlp docx_question_level）：
      1. 标题层级栈：遇 Heading N 维护父子栈（level<=栈顶时弹栈），flush「祖先路径 +
         正文」成一个块——比现状只按 H1 切粒度更细、上下文更全。
      2. 表格原子化：表格作为独立块，永不进文本切分 → 治 SentenceSplitter 把表格从
         | 处锯断的病（latest.json 实证「标称值 | 750 | 580 | 420」被截断）。

    融合项目优势：用 body XML children 遍历保留表格在段落流中的位置（RAGFlow 原版
    分别遍历 paragraphs/tables 会丢位置）；复用 serialize_table 的「列名: 值」序列化。
    """
    p = Path(file_path)
    try:
        from docx import Document as _DocxDocument
    except ImportError:
        logger.warning("python-docx 未安装，DOCX 结构化切块不可用: %s", p.name)
        return [], []

    try:
        d = _DocxDocument(str(p))
    except Exception as exc:
        logger.warning("打开 .docx 失败 %s: %s", p.name, exc)
        return [], []

    body = d.element.body
    # xml element id → table 对象，O(1) 查表（与 extract_docx_sections 同构）
    tbl_map: dict[int, object] = {id(t._element): t for t in d.tables}

    # RAGFlow Manual 标题栈：question_stack=标题文本，level_stack=对应层级
    question_stack: list[str] = []
    level_stack: list[int] = []
    last_answer_parts: list[str] = []  # 当前最近标题下累积的正文段落

    chunks: list[str] = []
    sections: list[str] = []

    def _current_section() -> str:
        return " > ".join(question_stack) if question_stack else ""

    def _flush_answer() -> None:
        """把当前累积正文 flush 成文本块（附当前祖先标题路径）；超长则内部二次切。"""
        if not last_answer_parts:
            return
        text = "\n\n".join(last_answer_parts)
        last_answer_parts.clear()
        if not text.strip():
            return
        section = _current_section()
        if len(text) > max_section_chars:
            for sub in _split_oversize(text, max_section_chars):
                chunks.append(sub)
                sections.append(section)
        else:
            chunks.append(text)
            sections.append(section)

    def _flush_table(tbl_obj) -> None:
        """表格作为独立原子块（不参与文本切分），section 归当前最近标题。"""
        serialized = serialize_table(tbl_obj)
        if serialized and serialized.strip():
            chunks.append(serialized)
            sections.append(_current_section())

    for child in list(body):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            level = _heading_level(child)
            txt = _paragraph_text(child)

            if level > 0 and txt:
                # 遇标题：先 flush 上一段正文，再维护标题栈
                _flush_answer()
                # RAGFlow 弹栈：弹出层级 >= 当前的栈顶（level <= top 时 pop），
                # 使新标题挂到比它浅的父标题下
                while level_stack and level <= level_stack[-1]:
                    level_stack.pop()
                    question_stack.pop()
                question_stack.append(txt)
                level_stack.append(level)
                continue

            # 非标题段落：图片占位 or 正文累积
            if _has_image(child):
                # 埋 [[IMG:sha16]] 占位符，ingestion 阶段按位置清单回填真实 VLM 描述。
                # Word 自带图注 txt 继续保留（权威命名）；占位符前置于图注。
                sha = _image_blob_sha(child, d)
                if sha:
                    last_answer_parts.append(
                        f"[[IMG:{sha[:16]}]]" + (f" {txt}" if txt else "")
                    )
                elif txt:
                    last_answer_parts.append(txt)
                # 无 sha 又无图注：不再制造 "[图: 电路图]" 噪声，直接跳过
            elif txt:
                last_answer_parts.append(txt)

        elif tag == "tbl":
            tbl_obj = tbl_map.get(id(child))
            if tbl_obj is not None:
                # 表格前先 flush 正文，保持「正文块、表格块」的文档原始顺序
                _flush_answer()
                _flush_table(tbl_obj)

    _flush_answer()

    logger.info("chunk_docx_structured: %s → %d chunks", p.name, len(chunks))
    return chunks, sections
