"""chunk_docx_structured 单测：移植自 RAGFlow Manual 的 DOCX 结构化切块。

重点验证两招：
  1. 表格原子化——整张表格落在单个 chunk，不被切开（治 SentenceSplitter 从 | 处锯断）。
  2. 标题层级栈——H2 下的内容带 "H1 > H2" 祖先路径；同级 H2 切换时栈正确回退。
"""
from __future__ import annotations

from pathlib import Path

from core.rag.chunking.ragflow_manual_docx import chunk_docx_structured


def _make_docx(path: Path) -> None:
    """构造一个含 H1/H2 + Word 原生表格的测试 DOCX。"""
    from docx import Document

    doc = Document()
    doc.add_heading("实验三 电阻串并联电路", level=1)
    doc.add_paragraph("本实验研究串并联电路的特性。")
    doc.add_heading("串联电路", level=2)
    doc.add_paragraph("按图接线，调R2=0，测量I、U、U1、U2。")

    # Word 原生表格：标称值/实测值/单位，模拟 latest.json 里被截断的那种表
    table = doc.add_table(rows=3, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "标称值", "实测值", "单位"
    r1 = table.rows[1].cells
    r1[0].text, r1[1].text, r1[2].text = "750", "745", "Ω"
    r2 = table.rows[2].cells
    r2[0].text, r2[1].text, r2[2].text = "580", "578", "Ω"

    doc.add_heading("并联电路", level=2)
    doc.add_paragraph("并联电路的总电阻计算。")
    doc.save(str(path))


def test_table_not_split(tmp_path: Path):
    """表格原子化：整表（含 750/580 多行）落在同一 chunk，不被切开。"""
    fp = tmp_path / "t.docx"
    _make_docx(fp)
    chunks, _ = chunk_docx_structured(str(fp), max_section_chars=1200)

    chunks_with_750 = [c for c in chunks if "750" in c]
    chunks_with_580 = [c for c in chunks if "580" in c]
    assert chunks_with_750, "表格首行数据 750 应出现在某 chunk"
    assert chunks_with_580, "表格次行数据 580 应出现在某 chunk"
    # 同一张表的两行必须在同一 chunk（原子化，不被 SentenceSplitter 锯断）
    assert chunks_with_750[0] is chunks_with_580[0], (
        "表格应原子化为单个 chunk，两行数据不得分散到不同 chunk"
    )
    # 表头也在该 chunk
    assert "标称值" in chunks_with_750[0] and "实测值" in chunks_with_750[0]


def test_heading_section_path(tmp_path: Path):
    """标题层级栈：H2 下内容带 "H1 > H2" 祖先路径，同级 H2 切换时栈回退正确。"""
    fp = tmp_path / "t.docx"
    _make_docx(fp)
    chunks, sections = chunk_docx_structured(str(fp), max_section_chars=1200)

    # H1 正文段 → section = "实验三 电阻串并联电路"
    h1_idx = next(i for i, c in enumerate(chunks) if "本实验研究串并联电路" in c)
    assert sections[h1_idx] == "实验三 电阻串并联电路"

    # H2「串联电路」正文 → section 带 H1 前缀
    h2_idx = next(i for i, c in enumerate(chunks) if "调R2=0" in c)
    assert sections[h2_idx] == "实验三 电阻串并联电路 > 串联电路"

    # 同级 H2「并联电路」→ 栈应回退到 H1 下，section 是 "H1 > 并联电路"
    # （而非 "H1 > 串联电路 > 并联电路"，验证弹栈）
    h2b_idx = next(i for i, c in enumerate(chunks) if "总电阻计算" in c)
    assert sections[h2b_idx] == "实验三 电阻串并联电路 > 并联电路"


def test_table_section_under_heading(tmp_path: Path):
    """表格块的 section 归其最近的标题（H2「串联电路」下）。"""
    fp = tmp_path / "t.docx"
    _make_docx(fp)
    chunks, sections = chunk_docx_structured(str(fp), max_section_chars=1200)

    table_idx = next(i for i, c in enumerate(chunks) if "750" in c)
    assert sections[table_idx] == "实验三 电阻串并联电路 > 串联电路"


def test_oversize_section_fallback(tmp_path: Path):
    """超长章节正文回退二次切：正文超 max_section_chars 切成多块，表格仍原子。"""
    from docx import Document

    fp = tmp_path / "big.docx"
    doc = Document()
    doc.add_heading("超长章节", level=1)
    doc.add_paragraph("内容段落。" * 1000)  # 远超阈值
    # 表格紧随其后，验证它不被超长切波及
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text, table.cell(0, 1).text = "表头A", "表头B"
    table.cell(1, 0).text, table.cell(1, 1).text = "数据1", "数据2"
    doc.save(str(fp))

    chunks, _ = chunk_docx_structured(str(fp), max_section_chars=500)
    # 超长正文应被切成多块（>1）
    text_chunks = [c for c in chunks if "内容段落" in c]
    assert len(text_chunks) > 1, "超长正文应回退二次切成多块"
    # 表格仍原子（单块，且含全部数据）
    table_chunks = [c for c in chunks if "数据1" in c and "数据2" in c]
    assert len(table_chunks) == 1, "表格即使前文超长，仍应原子化为单块"


def test_empty_docx_returns_empty(tmp_path: Path):
    """空 DOCX（无任何段落/表格）返回空列表，不报错。"""
    from docx import Document

    fp = tmp_path / "empty.docx"
    Document().save(str(fp))  # 空文档
    chunks, sections = chunk_docx_structured(str(fp), max_section_chars=1200)
    assert chunks == []
    assert sections == []
