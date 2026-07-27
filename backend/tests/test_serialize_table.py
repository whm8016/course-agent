"""serialize_table 单测：Markdown 序列化 + 合并单元格去重 + 单元格内换行。

锁定的三个关键行为（治 latest.json 实证的表格残缺）：
  1. 基本表格 → 规范 Markdown（| col | col | + --- 分隔符），数据完整不丢。
  2. 横向合并 → 合并值只出现一次（id(cell._tc) 去重），列对齐。
  3. 单元格内换行 → 规范为空格，不破坏 Markdown 行结构（治「标称值\\n实测值」断行）。
"""
from __future__ import annotations


def _serialize(table) -> str:
    from core.rag.llamaindex.file_routing import serialize_table

    return serialize_table(table)


def test_serialize_table_basic_markdown():
    """基本测量表 → 规范 Markdown，数据完整。"""
    from docx import Document

    doc = Document()
    table = doc.add_table(rows=3, cols=3)
    for c, h in enumerate(["标称值", "实测值", "单位"]):
        table.cell(0, c).text = h
    for r, row in enumerate([("750", "745", "Ω"), ("580", "578", "Ω")], 1):
        for c, v in enumerate(row):
            table.cell(r, c).text = v

    out = _serialize(table)
    assert "| 标称值 | 实测值 | 单位 |" in out
    assert "| 750 | 745 | Ω |" in out
    assert "| 580 | 578 | Ω |" in out
    # Markdown 表头分隔符
    assert "---" in out


def test_serialize_table_horizontal_merge():
    """横向合并：合并值只出现一次（不被重复引用复制），其余数据不丢。"""
    from docx import Document

    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    for r in range(2):
        for c in range(3):
            table.cell(r, c).text = f"r{r}c{c}"
    # 横向合并第 0 行前两格
    table.cell(0, 0).merge(table.cell(0, 1)).text = "HSPAN"

    out = _serialize(table)
    # 合并值只出现一次（python-docx row.cells 对横向合并返回重复引用，按 id(_tc) 去重）
    assert out.count("HSPAN") == 1
    # 未被合并波及的格子仍在
    assert "r0c2" in out
    assert "r1c0" in out
    # Markdown 格式
    assert out.startswith("| ")


def test_serialize_table_intra_cell_newline():
    """单元格内换行规范为空格，不把 Markdown 行从中间断开。

    复现 latest.json 实证：原表左上格「标称值\\n实测值」带换行，旧实现把 Markdown
    表头行断成两行。修复后换行→空格，表头保持单行完整。
    """
    from docx import Document

    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "标称值\n实测值"  # 单元格内换行
    table.cell(0, 1).text = "数据"
    table.cell(1, 0).text = "R1"
    table.cell(1, 1).text = "100"

    out = _serialize(table)
    # 以 | 开头的行 = Markdown 表格行
    md_lines = [ln for ln in out.splitlines() if ln.startswith("|")]
    assert md_lines, "应输出 Markdown 表格行"
    # 表头行（第一行）应含规范后的「标称值 实测值」，且与「数据」在同一行
    assert "标称值 实测值" in md_lines[0]
    assert "数据" in md_lines[0]
    # 不应出现换行把「实测值 | 数据」拆到不同行（旧 bug 的特征）
    assert md_lines[0].endswith("|")
