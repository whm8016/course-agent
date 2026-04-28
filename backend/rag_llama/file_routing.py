"""
File Type Router
================

Centralized file type classification and routing for the RAG pipeline.
Determines the appropriate processing method for each document type.
"""

from dataclasses import dataclass
from enum import Enum
from pathlib import Path
from typing import List

import logging 
logger = logging.getLogger("FileTypeRouter")




class DocumentType(Enum):
    """Document type classification."""

    PDF = "pdf"
    TEXT = "text"
    MARKDOWN = "markdown"
    DOCX = "docx"
    IMAGE = "image"
    UNKNOWN = "unknown"


@dataclass
class FileClassification:
    """Result of file classification."""

    parser_files: List[str]
    text_files: List[str]
    docx_files: List[str]
    unsupported: List[str]


class FileTypeRouter:
    """File type router for the RAG pipeline.

    Classifies files before processing to route them to appropriate handlers:
    - PDF files -> PDF parsing
    - Text files -> Direct read (fast, simple)
    - Unsupported -> Skip with warning
    """

    PARSER_EXTENSIONS = {".pdf"}

    TEXT_EXTENSIONS = {
        ".txt",
        ".text",
        ".log",
        ".md",
        ".markdown",
        ".rst",
        ".asciidoc",
        ".json",
        ".yaml",
        ".yml",
        ".toml",
        ".csv",
        ".tsv",
        ".tex",
        ".latex",
        ".bib",
        ".py",
        ".js",
        ".ts",
        ".jsx",
        ".tsx",
        ".java",
        ".c",
        ".cpp",
        ".h",
        ".hpp",
        ".go",
        ".rs",
        ".rb",
        ".php",
        ".swift",
        ".kt",
        ".scala",
        ".r",
        ".sql",
        ".sh",
        ".bash",
        ".zsh",
        ".ps1",
        ".html",
        ".htm",
        ".xml",
        ".css",
        ".scss",
        ".sass",
        ".less",
        ".ini",
        ".cfg",
        ".conf",
        ".env",
        ".properties",
    }

    DOCX_EXTENSIONS = {".docx", ".doc"}
    IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".tiff", ".tif"}

    @classmethod
    def get_document_type(cls, file_path: str) -> DocumentType:
        """Classify a single file by its type."""
        ext = Path(file_path).suffix.lower()

        if ext in cls.PARSER_EXTENSIONS:
            return DocumentType.PDF
        elif ext in cls.TEXT_EXTENSIONS:
            return DocumentType.TEXT
        elif ext in cls.DOCX_EXTENSIONS:
            return DocumentType.DOCX
        elif ext in cls.IMAGE_EXTENSIONS:
            return DocumentType.IMAGE
        else:
            if cls._is_text_file(file_path):
                return DocumentType.TEXT
            return DocumentType.UNKNOWN

    @classmethod
    def _is_text_file(cls, file_path: str, sample_size: int = 8192) -> bool:
        """Detect if a file is text-based by examining its content."""
        try:
            with open(file_path, "rb") as f:
                chunk = f.read(sample_size)

            if b"\x00" in chunk:
                return False

            chunk.decode("utf-8")
            return True
        except (UnicodeDecodeError, IOError, OSError):
            return False

    @classmethod
    def classify_files(cls, file_paths: List[str]) -> FileClassification:
        """Classify a list of files by processing method."""
        parser_files = []
        text_files = []
        docx_files = []
        unsupported = []

        for path in file_paths:
            doc_type = cls.get_document_type(path)

            if doc_type == DocumentType.PDF:
                parser_files.append(path)
            elif doc_type in (DocumentType.TEXT, DocumentType.MARKDOWN):
                text_files.append(path)
            elif doc_type == DocumentType.DOCX:
                ext = Path(path).suffix.lower()
                if ext == ".docx":
                    docx_files.append(path)
                else:
                    # legacy binary .doc — not supported by python-docx; keep explicit
                    unsupported.append(path)
            else:
                unsupported.append(path)

        logger.debug(
            f"Classified {len(file_paths)} files: "
            f"{len(parser_files)} parser, {len(text_files)} text, {len(docx_files)} docx, "
            f"{len(unsupported)} unsupported"
        )

        return FileClassification(
            parser_files=parser_files,
            text_files=text_files,
            docx_files=docx_files,
            unsupported=unsupported,
        )

    @classmethod
    def extract_docx_text(cls, file_path: str) -> str:
        """Read plain text from a .docx (Office Open XML). Legacy .doc is not supported."""
        p = Path(file_path)
        if p.suffix.lower() != ".docx":
            return ""
        try:
            from docx import Document
        except ImportError:
            logger.warning("python-docx is not installed; cannot read .docx")
            return ""
        try:
            d = Document(str(p))
            parts: list[str] = []
            for para in d.paragraphs:
                t = (para.text or "").strip()
                if t:
                    parts.append(t)
            for table in d.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n\n".join(parts)
        except Exception as exc:
            logger.warning("Failed to read .docx %s: %s", p.name, exc)
            return ""

    @classmethod
    def extract_docx_sections(cls, file_path: str) -> list[dict]:
        """Parse a .docx and split it into sections by Heading 1.

        Each returned dict has:
          - title   : heading text (e.g. "实验九 二端口网络研究")
          - content : full text of that section, including table cells with
                      column-header prefixes and [图: 电路图] image placeholders
          - metadata: {"section": title, "file_name": basename}

        Falls back to a single section containing the whole document when
        no Heading 1 paragraphs are found (e.g. plain text docs).
        """
        p = Path(file_path)
        if p.suffix.lower() != ".docx":
            return []
        try:
            from docx import Document as _DocxDocument
        except ImportError:
            logger.warning("python-docx is not installed; cannot read .docx")
            return []

        NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        DML = "http://schemas.openxmlformats.org/drawingml/2006/main"
        PIC = "http://schemas.openxmlformats.org/drawingml/2006/picture"

        try:
            d = _DocxDocument(str(p))
        except Exception as exc:
            logger.warning("Failed to open .docx %s: %s", p.name, exc)
            return []

        # Build a lookup: docx table objects → their serialized text
        def _serialize_table(tbl) -> str:
            rows = tbl.rows
            if not rows:
                return ""
            # Treat first row as header if it contains non-empty cells
            first_cells = [c.text.strip() for c in rows[0].cells]
            has_header = any(first_cells)
            headers = first_cells if has_header else []
            lines: list[str] = []
            start = 1 if has_header else 0
            if has_header:
                lines.append(" | ".join(h for h in headers if h))
            for row in rows[start:]:
                cells = [c.text.strip() for c in row.cells]
                if not any(cells):
                    continue
                if headers:
                    pairs = []
                    for h, v in zip(headers, cells):
                        if h and v:
                            pairs.append(f"{h}: {v}")
                        elif v:
                            pairs.append(v)
                    lines.append(" | ".join(pairs))
                else:
                    lines.append(" | ".join(c for c in cells if c))
            return "\n".join(lines)

        # Walk body XML children to preserve paragraph/table order per section
        body = d.element.body
        sections: list[dict] = []
        current_title: str | None = None
        current_parts: list[str] = []

        # Map xml element id → table object for quick lookup
        tbl_map: dict[int, object] = {id(t._element): t for t in d.tables}

        def _flush(title: str | None, parts: list[str]) -> None:
            if parts:
                content = "\n\n".join(parts)
                t = title or ""
                sections.append({
                    "title": t,
                    "content": (f"{t}\n\n{content}").strip() if t else content,
                    "metadata": {"section": t, "file_name": p.name},
                })

        # pStyle val values that map to Heading 1 across different Word versions/locales:
        # "Heading1"  — English style id
        # "1"         — common numeric alias
        # "2"         — OOXML built-in numeric id observed in Chinese Word installs
        _H1_SVALS = {"Heading1", "1", "2"}

        for child in list(body):
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if tag == "p":
                pStyle = child.find(f".//{{{NS}}}pStyle")
                sval = (pStyle.get(f"{{{NS}}}val") or "") if pStyle is not None else ""
                txt = "".join(t.text or "" for t in child.iter(f"{{{NS}}}t")).strip()

                # Heading 1 → start new section (require non-empty title)
                if sval in _H1_SVALS and txt:
                    _flush(current_title, current_parts)
                    current_title = txt
                    current_parts = []
                    continue

                # Image detection
                has_image = (
                    child.find(f".//{{{DML}}}blip") is not None
                    or child.find(f".//{{{PIC}}}pic") is not None
                )
                if has_image:
                    label = f"[图: {txt}]" if txt else "[图: 电路图]"
                    current_parts.append(label)
                elif txt:
                    current_parts.append(txt)

            elif tag == "tbl":
                tbl_obj = tbl_map.get(id(child))
                if tbl_obj is not None:
                    serialized = _serialize_table(tbl_obj)
                    if serialized:
                        current_parts.append(serialized)

        _flush(current_title, current_parts)

        # If no Heading 1 was found, fall back to whole-doc single section
        if not sections:
            fallback = cls.extract_docx_text(file_path)
            if fallback:
                sections.append({
                    "title": p.stem,
                    "content": fallback,
                    "metadata": {"section": p.stem, "file_name": p.name},
                })

        logger.info(
            "extract_docx_sections: %s → %d sections", p.name, len(sections)
        )
        return sections

    @classmethod
    def read_text_file_sync(cls, file_path: str) -> str:
        """Read a text file with automatic encoding detection (sync; 与 LightRAG 线程池解析共用)."""
        encodings = ["utf-8", "utf-8-sig", "gbk", "gb2312", "gb18030", "latin-1", "cp1252"]

        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue

        with open(file_path, "rb") as f:
            return f.read().decode("utf-8", errors="replace")

    @classmethod
    def extract_pdf_text(cls, file_path: str) -> str:
        """Extract PDF text with PyMuPDF（LlamaIndex / LightRAG 摄入共用）。"""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            logger.warning("PyMuPDF not installed. Cannot extract PDF text.")
            return ""
        try:
            doc = fitz.open(file_path)
            texts = []
            for page in doc:
                texts.append(page.get_text())
            doc.close()
            return "\n\n".join(texts)
        except Exception as exc:
            logger.warning("Failed to extract PDF text %s: %s", file_path, exc)
            return ""

    @classmethod
    async def read_text_file(cls, file_path: str) -> str:
        """Read a text file with automatic encoding detection."""
        return cls.read_text_file_sync(file_path)

    @classmethod
    def needs_parser(cls, file_path: str) -> bool:
        """Quick check if a single file needs parser processing."""
        doc_type = cls.get_document_type(file_path)
        return doc_type in (DocumentType.PDF, DocumentType.DOCX, DocumentType.IMAGE)

    @classmethod
    def is_text_readable(cls, file_path: str) -> bool:
        """Check if a file can be read directly as text."""
        doc_type = cls.get_document_type(file_path)
        return doc_type in (DocumentType.TEXT, DocumentType.MARKDOWN)

    @classmethod
    def get_supported_extensions(cls) -> set[str]:
        """Get the set of all supported file extensions."""
        return cls.PARSER_EXTENSIONS | cls.TEXT_EXTENSIONS

    @classmethod
    def get_glob_patterns(cls) -> list[str]:
        """Get glob patterns for file searching."""
        return [f"*{ext}" for ext in sorted(cls.get_supported_extensions())]
